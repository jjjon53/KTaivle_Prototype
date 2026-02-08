from docx import Document
from docx.shared import Pt
import os


class INDDocxGenerator:
    def __init__(self):
        self.doc = Document()
        self._set_styles()

    def _set_styles(self):
        # Set default font to Times New Roman 12pt
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

    def add_content(self, text: str):
        lines = text.split("\n")
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("## "):
                self.doc.add_heading(stripped[3:], level=1)
            elif stripped.startswith("### "):
                self.doc.add_heading(stripped[4:], level=2)
            elif stripped.startswith("#### "):
                self.doc.add_heading(stripped[5:], level=3)
            elif stripped:
                self.doc.add_paragraph(stripped)

    def add_disclaimer(self):
        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        run = p.add_run(
            "DISCLAIMER: This document is an AI-generated draft for prototype demonstration purposes only. It does NOT replace experimental data or professional regulatory review. Verify all values against original study reports."
        )
        run.bold = True
        run.font.color.rgb = None  # Default color or red? Keep simple.

    def save(self, output_path: str):
        self.doc.save(output_path)


def generate_docx(content: str, output_path: str):
    generator = INDDocxGenerator()
    generator.add_content(content)
    generator.add_disclaimer()
    generator.save(output_path)
